import os
import cv2
import numpy as np
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import tempfile
import openpyxl
from openpyxl.drawing.image import Image as XLImage
import sys

class AdicionarMarcaDagua:
    def __init__(self, marca_dagua_path, opacidade=0.3):
        """
        Inicializa o objeto para adicionar marca d'água
        :param marca_dagua_path: Caminho para a imagem da marca d'água
        :param opacidade: Valor de 0 a 1 para a opacidade da marca
        """
        self.marca_dagua = Image.open(marca_dagua_path).convert('RGBA')
        self.opacidade = opacidade
        self.marca_dagua_path = marca_dagua_path
        
    def redimensionar_marca(self, tamanho_desejado):
        """
        Redimensiona a marca d'água mantendo a proporção
        """
        ratio = min(tamanho_desejado[0]/self.marca_dagua.size[0], 
                   tamanho_desejado[1]/self.marca_dagua.size[1])
        novo_tamanho = (int(self.marca_dagua.size[0] * ratio),
                       int(self.marca_dagua.size[1] * ratio))
        return self.marca_dagua.resize(novo_tamanho, Image.LANCZOS)
        
    def adicionar_marca_imagem(self, imagem_path, output_path):
        """Adiciona marca d'água em uma imagem"""
        # Abrir a imagem
        img = Image.open(imagem_path).convert('RGBA')
        
        # Redimensionar marca d'água para 30% do tamanho da imagem
        tamanho_desejado = (int(img.size[0] * 0.3), int(img.size[1] * 0.3))
        marca_redimensionada = self.redimensionar_marca(tamanho_desejado)
        
        # Criar uma nova imagem com a marca d'água
        watermark = Image.new('RGBA', img.size, (0,0,0,0))
        
        # Calcular posição para centralizar
        x = (img.size[0] - marca_redimensionada.size[0]) // 2
        y = (img.size[1] - marca_redimensionada.size[1]) // 2
        
        # Ajustar opacidade da marca d'água
        marca_array = np.array(marca_redimensionada)
        marca_array[:, :, 3] = marca_array[:, :, 3] * self.opacidade
        marca_redimensionada = Image.fromarray(marca_array)
        
        # Colar a marca d'água
        watermark.paste(marca_redimensionada, (x, y), marca_redimensionada)
        
        # Combinar imagens
        combined = Image.alpha_composite(img, watermark)
        combined.convert('RGB').save(output_path)
        
    def adicionar_marca_pdf(self, pdf_path, output_path):
        """Adiciona marca d'água em um PDF"""
        # Converter PDF para imagens
        images = convert_from_path(pdf_path)
        
        # Processar cada página
        processed_images = []
        for img in images:
            # Converter para RGBA
            img_rgba = img.convert('RGBA')
            
            # Redimensionar marca d'água para 30% do tamanho da página
            tamanho_desejado = (int(img_rgba.size[0] * 0.3), int(img_rgba.size[1] * 0.3))
            marca_redimensionada = self.redimensionar_marca(tamanho_desejado)
            
            # Criar uma nova imagem com a marca d'água
            watermark = Image.new('RGBA', img_rgba.size, (0,0,0,0))
            
            # Calcular posição para centralizar
            x = (img_rgba.size[0] - marca_redimensionada.size[0]) // 2
            y = (img_rgba.size[1] - marca_redimensionada.size[1]) // 2
            
            # Ajustar opacidade da marca d'água
            marca_array = np.array(marca_redimensionada)
            marca_array[:, :, 3] = marca_array[:, :, 3] * self.opacidade
            marca_redimensionada = Image.fromarray(marca_array)
            
            # Colar a marca d'água
            watermark.paste(marca_redimensionada, (x, y), marca_redimensionada)
            
            # Combinar imagens
            combined = Image.alpha_composite(img_rgba, watermark)
            processed_images.append(combined.convert('RGB'))
        
        # Salvar PDF
        processed_images[0].save(output_path, save_all=True, 
                               append_images=processed_images[1:])
        
    def adicionar_marca_doc(self, doc_path, output_path):
        """Adiciona marca d'água em um documento Word"""
        doc = Document(doc_path)
        
        # Redimensionar marca d'água para um tamanho adequado para o documento
        tamanho_desejado = (600, 400)  # Tamanho base para documentos
        marca_redimensionada = self.redimensionar_marca(tamanho_desejado)
        
        # Ajustar opacidade
        marca_array = np.array(marca_redimensionada)
        marca_array[:, :, 3] = marca_array[:, :, 3] * self.opacidade
        marca_redimensionada = Image.fromarray(marca_array)
        
        # Salvar temporariamente
        temp_img_path = "temp_watermark.png"
        marca_redimensionada.save(temp_img_path)
        
        # Adicionar marca d'água em cada seção
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(temp_img_path, width=Inches(8))
        
        # Salvar documento
        doc.save(output_path)
        
        # Remover arquivo temporário
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)

    def adicionar_marca_excel(self, excel_path, output_path):
        """Adiciona marca d'água em uma planilha Excel"""
        # Carregar o arquivo Excel
        workbook = openpyxl.load_workbook(excel_path)
        
        # Redimensionar marca d'água
        tamanho_desejado = (300, 200)  # Tamanho base para Excel
        marca_redimensionada = self.redimensionar_marca(tamanho_desejado)
        
        # Ajustar opacidade
        marca_array = np.array(marca_redimensionada)
        marca_array[:, :, 3] = marca_array[:, :, 3] * self.opacidade
        marca_redimensionada = Image.fromarray(marca_array)
        
        # Salvar temporariamente
        temp_img_path = "temp_watermark.png"
        marca_redimensionada.save(temp_img_path)
        
        # Adicionar marca d'água em cada planilha
        for sheet in workbook.sheetnames:
            ws = workbook[sheet]
            
            # Carregar a imagem
            img = XLImage(temp_img_path)
            
            # Calcular posição central (célula B2 como padrão)
            ws.add_image(img, 'B2')
        
        # Salvar arquivo
        workbook.save(output_path)
        
        # Remover arquivo temporário
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)

def main():
    # Verificar argumentos da linha de comando
    if len(sys.argv) != 4:
        print("Uso: python script.py arquivo_entrada arquivo_saida marca_dagua")
        sys.exit(1)
        
    arquivo_entrada = sys.argv[1]
    arquivo_saida = sys.argv[2]
    marca_dagua = sys.argv[3]
    
    # Instanciar o processador
    processador = AdicionarMarcaDagua(
        marca_dagua_path=marca_dagua,
        opacidade=0.3
    )
    
    # Processar o arquivo baseado na extensão
    extensao = os.path.splitext(arquivo_entrada)[1].lower()
    
    if extensao in ['.jpg', '.jpeg', '.png']:
        processador.adicionar_marca_imagem(arquivo_entrada, arquivo_saida)
    elif extensao == '.pdf':
        processador.adicionar_marca_pdf(arquivo_entrada, arquivo_saida)
    elif extensao in ['.doc', '.docx']:
        processador.adicionar_marca_doc(arquivo_entrada, arquivo_saida)
    elif extensao == '.xlsx':
        processador.adicionar_marca_excel(arquivo_entrada, arquivo_saida)
    else:
        print(f"Formato de arquivo não suportado: {extensao}")
        sys.exit(1)

if __name__ == "__main__":
    main() 